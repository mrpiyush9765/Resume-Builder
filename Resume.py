
import streamlit as st
from docx import Document
from PyPDF2 import PdfReader
from io import BytesIO

st.set_page_config(page_title="AI Resume Builder", layout="centered")

st.title("🤖 Smart AI Resume Builder")
st.markdown("Fill in details or upload your resume to get started!")

# 1. Template Selection
template = st.selectbox("Choose a Resume Template", ["Modern Basic", "Bold Header", "Minimalist Professional"])

# 2. Option to Upload Existing Resume
uploaded_file = st.file_uploader("📄 Upload Existing Resume (PDF)", type=["pdf"])
if uploaded_file:
    reader = PdfReader(uploaded_file)
    extracted_text = ""
    for page in reader.pages:
        extracted_text += page.extract_text() + "\n"
    st.text_area("📜 Extracted Resume Text", value=extracted_text, height=300)

st.divider()

# 3. Resume Input Fields
st.header("📋 Enter Your Resume Details")
name = st.text_input("Full Name")
email = st.text_input("Email")
phone = st.text_input("Phone Number")
summary = st.text_area("Professional Summary")
skills = st.text_area("Skills (comma-separated)")
education = st.text_area("Education")
experience = st.text_area("Work Experience")
certifications = st.text_area("Certifications")
projects = st.text_area("Projects")
Language = st.text_area("Language")

if st.button("Generate Resume"):
    doc = Document()

    # Apply selected template style (basic examples)
    if template == "Modern Basic":
        doc.add_heading(name, 0)
    elif template == "Bold Header":
        doc.add_heading(name.upper(), 0)
    else:
        doc.add_paragraph(f"Name: {name}", style='List Bullet')

    doc.add_paragraph(f"Email: {email} | Phone: {phone}")
    doc.add_paragraph("\n")

    doc.add_heading("Professional Summary", level=1)
    doc.add_paragraph(summary)

    doc.add_heading("Skills", level=1)
    doc.add_paragraph(skills)

    doc.add_heading("Education", level=1)
    doc.add_paragraph(education)

    doc.add_heading("Work Experience", level=1)
    doc.add_paragraph(experience)

    doc.add_heading("Projects", level=1)
    doc.add_paragraph(projects)

    doc.add_heading("Certifications", level=1)
    doc.add_paragraph(certifications)

    doc.add_heading("Language", level=1)
    doc.add_paragraph(certifications)

    buffer = BytesIO()
    doc.save(buffer)
    buffer.seek(0)

    st.success("✅ Resume generated successfully!")
    st.download_button("📄 Download Resume (.docx)", data=buffer, file_name=f"{name.replace(' ', '_')}_Resume.docx",
                       mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")


